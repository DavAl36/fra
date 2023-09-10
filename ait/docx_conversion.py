import docx
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH



# Create an instance of a word document
def txt_to_docx(path_in, path_out):

    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading('Book', 0)

    # open the data file
    file = open(path_in, mode="r", encoding="utf-8")
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT #RIGHT/CENTER
            

    for line in file.readlines():
        #if ('PHOTO' in line):
        #        doc.add_heading(line, 3)
        if(len(line) > 2):
            if('PHOTO' in line[1:7]):
                para = p.add_run(line)
                para.font.name = 'Arial'
                para.font.size = Pt(8) 
                para.font.bold = True
                para.font.italic = True
            else:
                para = p.add_run(line)
                para.font.name = 'Arial'
                para.font.size = Pt(12)    
                para.font.bold = False
                para.font.italic = False



    '''
    text = ''
    for line in file.readlines():
        #if ('PHOTO' in line):
        #        doc.add_heading(line, 3)
        if(len(line) > 2):
            text = text + line
    
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT #RIGHT/CENTER
    para = p.add_run(text)
    para.font.name = 'Arial'
    para.font.size = Pt(12)
    '''

    file.close()
    doc.save(path_out)
    os.chmod(path_out,0o777)





